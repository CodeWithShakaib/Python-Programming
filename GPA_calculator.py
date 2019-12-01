def username():
    import getpass
    username = getpass.getuser()
    return username


import os

def clear():
    os.system("CLS")
os.system("color F0")
c_h_List=[]
q_p_List=[]
o_m_List=[]
t_m_List=[]
per_list=[]
grade_list=[]
ob_m_List=[]
obtained_marks_list=[]
subject_name=[]
sr_list=[]
Q_P_1={
    1 : 0.00,
    2 : 0.00,
    3 : 0.00,
    4 : 0.00,
    5 : 0.00,
    6 : 0.00,
    7 : 0.00,
    8 : 1.00,
    10 : 2.00,
    13 : 3.00,
    16 : 4.00,
    17 : 4.00,
    18 : 4.00,
    19 : 4.00,
    20 : 4.00
}
Q_P_2={
    1 : 0.00,
    2 : 0.00,
    3 : 0.00,
    4 : 0.00,
    5 : 0.00,
    6 : 0.00,
    7 : 0.00,
    8 : 0.00,
    9 : 0.00,
    10 : 0.00,
    11 : 0.00,
    12 : 0.00,
    13 : 0.00,
    14 : 0.00,
    15 : 0.00,
    16 : 2.00,
    17 : 2.60,
    18 : 3.00,
    19 : 3.60,
    20 : 4.00,
    21 : 4.40,
    22 : 4.60,
    23 : 5.00,
    24 : 5.40,
    25 : 5.60,
    26 : 6.00,
    27 : 6.40,
    28 : 6.60,
    29 : 7.00,
    30 : 7.40,
    31 : 7.60,
    32 : 8.00,
    33 : 8.00,
    34 : 8.00,
    35 : 8.00,
    36 : 8.00,
    37 : 8.00,
    38 : 8.00,
    39 : 8.00,
    40 : 8.00
}
Q_P_3={
    1 : 0.00,
    2 : 0.00,
    3 : 0.00,
    4 : 0.00,
    5 : 0.00,
    6 : 0.00,
    7 : 0.00,
    8 : 0.00,
    9 : 0.00,
    10 : 0.00,
    11 : 0.00,
    12 : 0.00,
    13 : 0.00,
    14 : 0.00,
    15 : 0.00,
    16 : 0.00,
    17 : 0.00,
    18 : 0.00,
    19 : 0.00,
    20 : 0.00,
    21 : 0.00,
    22 : 0.00,
    23 : 0.00,
    24 : 3.00,
    25 : 3.60,
    26 : 3.90,
    27 : 4.50,
    28 : 5.10,
    29 : 5.40,
    30 : 6.00,
    31 : 6.30,
    32 : 6.60,
    33 : 6.90,
    34 : 7.20,
    35 : 7.50,
    36 : 8.10,
    37 : 8.40,
    38 : 8.70,
    39 : 9.00,
    40 : 9.30,
    41 : 9.60,
    42 : 9.90,
    43 : 10.20,
    44 : 10.50,
    45 : 11.10,
    46 : 11.40,
    47 : 11.70,
    48 : 12.00,
    49 : 12.00,
    50 : 12.00,
    51 : 12.00,
    52 : 12.00,
    53 : 12.00,
    54 : 12.00,
    55 : 12.00,
    56 : 12.00,
    57 : 12.00,
    58 : 12.00,
    59 : 12.00,
    60 : 12.00
}
Q_P_4={
     1 : 0.00,
     2 : 0.00,
     3 : 0.00,
     4 : 0.00,
     5 : 0.00,
     6 : 0.00,
     7 : 0.00,
     8 : 0.00,
     9 : 0.00,
     10 : 0.00,
     11 : 0.00,
     12 : 0.00,
     13 : 0.00,
     14 : 0.00,
     15 : 0.00,
     16 : 0.00,
     17 : 0.00,
     18 : 0.00,
     19 : 0.00,
     20 : 0.00,
     21 : 0.00,
     22 : 0.00,
     23 : 0.00,
     24 : 0.00,
     25 : 0.00,
     26 : 0.00,
     27 : 0.00,
     28 : 0.00,
     29 : 0.00,
     30 : 0.00,
     31 : 0.00,
     32 : 4.00,
     33 : 4.40,
     34 : 5.20,
     35 : 5.60,
     36 : 6.00,
     37 : 6.40,
     38 : 7.20,
     39 : 7.60,
     40 : 8.00,
     41 : 8.40,
     42 : 8.80,
     43 : 8.80,
     44 : 9.20,
     45 : 9.60,
     46 : 10.00,
     47 : 10.40,
     48 : 10.80,
     49 : 10.80,
     50 : 11.20,
     51 : 11.60,
     52 : 12.00,
     53 : 12.00,
     54 : 12.40,
     55 : 12.80,
     56 : 13.20,
     57 : 13.60,
     58 : 14.00,
     59 : 14.40,
     60 : 14.80,
     61 : 14.80,
     62 : 15.20,
     63 : 15.60,
     64 : 16.00,
     65 : 16.00,
     66 : 16.00,
     67 : 16.00,
     68 : 16.00,
     69 : 16.00,
     70 : 16.00,
     71 : 16.00,
     72 : 16.00,
     73 : 16.00,
     74 : 16.00,
     75 : 16.00,
     76 : 16.00,
     77 : 16.00,
     78 : 16.00,
     79 : 16.00,
     80 : 16.00  
}
C_H={
    1 : 20,
    2 : 40,
    3 : 60,
    4 : 80,
    5 : 100
}
def personal_info():
    clear()
    print("___________________________________________________________________________________________________________________\n")
    print("\nPERSONAL DETAILS :\n")
    universty=input("  Enter universty Name ?")
    department=input("  Enter Departnment ?")
    name=input("  Enter Your Name ?")
    reg_no=input("  Enter Registration # ?")
    degree=input("  Enter Degree ?")
    section=input("  Enter Section ?")
    semyster=input("  Enter Semyster ?")
    print("___________________________________________________________________________________________________________________\n")
    return universty,department,name,reg_no,degree,section,semyster


def result_calculator():
    print("___________________________________________________________________________________________________________________\n")
    print("\n\nACEDIMIC DETAILS :\n")
    T_Q_P=0
    T_C_H=0
    b=0
    subjects=int(input("  Total Numbers Of Subjects You Are Studing ?"))
    print("___________________________________________________________________________________________________________________\n")
    for i in range(0,subjects):
        b+=1
        sr_list.append(b)
        sub_name=input("  Enter Subject Name ?")
        subject_name.append(sub_name)
        while True:
            credit_hours=int(input("  Enter Credit Hours ?"))
            if credit_hours==1 or credit_hours==2 or credit_hours==3 or credit_hours==4 or credit_hours==5:
                break
            else:
                print("You entered wrong credit hour ! please enter again")
        while True:
            obtained_marks=int(input("  Enter Obtainer Marks ?"))
            if credit_hours==1 and obtained_marks>=1 and obtained_marks<=7:
                ob_m_List.append(obtained_marks)
                break    
            elif credit_hours==2 and obtained_marks>=1 and obtained_marks<=16:
                ob_m_List.append(obtained_marks)
                break
            elif credit_hours==3 and obtained_marks>=1 and obtained_marks<=24:
                ob_m_List.append(obtained_marks)
                break
            elif credit_hours==4 and obtained_marks>=1 and obtained_marks<=32:
                ob_m_List.append(obtained_marks)
                break
            elif credit_hours==5 and obtained_marks>=1 and obtained_marks<=40:
                o_m_List.append(obtained_marks)
                break
            elif credit_hours==1 and obtained_marks>=8 and obtained_marks<=20:
                o_m_List.append(obtained_marks)
                break
            elif credit_hours==2 and obtained_marks>=16 and obtained_marks<=40:
                o_m_List.append(obtained_marks)
                break
            elif credit_hours==3 and obtained_marks>=24 and obtained_marks<=60:
                o_m_List.append(obtained_marks)
                break
            elif credit_hours==4 and obtained_marks>=32 and obtained_marks<=80:
                o_m_List.append(obtained_marks)
                break
            elif credit_hours==5 and obtained_marks>=40 and obtained_marks<=100:
                o_m_List.append(obtained_marks)
                break
            else:
                print("you enter wrong obtained marks!!! please entered again")
        obtained_marks_list.append(obtained_marks)    

        if credit_hours==1:
            q_p_List.append(Q_P_1[obtained_marks])
        elif credit_hours==2:
            q_p_List.append(Q_P_2[obtained_marks])
        elif credit_hours==3:
            q_p_List.append(Q_P_3[obtained_marks])
        elif credit_hours==4:
            q_p_List.append(Q_P_4[obtained_marks])
        else:
            print("You Entered Wrong Credit Hours")
            break
        p=(obtained_marks/C_H[credit_hours])*100
        p=round(p,3)
        per_list.append(p)
        c_h_List.append(credit_hours)
        t_m_List.append(C_H[credit_hours])
        T_Q_P+=q_p_List[i]
        T_C_H+=c_h_List[i]
        if p>=80:
            grade_list.append('A')
        elif p>=65:
            grade_list.append('B')
        elif p>=50:
            grade_list.append('C')
        elif p>=40:
            grade_list.append('D')
        else:
            grade_list.append('F')
        print("___________________________________________________________________________________________________________________\n")
    #print("sr#             :",sr_list)
    #print("subjects        : ",subject_name)    
    #print("credit hour     : ",c_h_List)
    #print("total marks     :",t_m_List)
    #print("obtaines marks  :",obtained_marks_list)
    #print("percentage list :",per_list)
    #print("grade list      :",grade_list)

    GPA=T_Q_P/T_C_H
    total_marks=sum(t_m_List)
    obtained_marks=sum(obtained_marks_list)
    percentage=(obtained_marks/total_marks)*100
    percentage=round(percentage,3)
    if percentage>=80:
            grade='A'
            status='EXCELLENT'
    elif percentage>=65:
            grade='B'
            status='GOOD'
    elif percentage>=50:
            grade='C'
            status='SATISFACTORY'
    elif percentage>=40:
            grade='D'
            status='NEED IMPROVMENT'
    else:
            grade='E'
            status='FAIL'
    return GPA , total_marks , obtained_marks , percentage , grade , status , subjects


def create_table(sr_list,subject_name,c_h_List,t_m_List,obtained_marks_list,per_list,grade_list,GPA,T_M,O_M,pe,G,S):
    from beautifultable import BeautifulTable
    t = BeautifulTable()
    t1 = BeautifulTable()
    t.append_column('Sr#',sr_list)
    t.append_column('COURSES',subject_name)
    t.append_column('CREDIT HOURS',c_h_List)
    t.append_column('TOTAL MARKS',t_m_List)
    t.append_column('OBTAINED MARKS',obtained_marks_list)
    t.append_column('%',per_list)
    t.append_column('GRADE',grade_list)
    print(t)
    print("\n\t")
    c_1=['GPA','TOTAL MARKS','OBTAINED MARKS']
    c_2=[GPA,T_M,O_M]
    c_3=['GRADE','PERSENTAGE','STATUS']
    c_4=[G,p,S]
    t1.append_column('',c_1)
    t1.append_column('',c_2)
    t1.append_column('',c_3)
    t1.append_column('',c_4)
    print(t1)

def create_world_doc(universty,department,name,reg_no,degree,section,semyster,sr_list,subject_name,c_h_List,t_m_List,obtained_marks_list,per_list,grade_list,GPA,T_M,O_M,pe,G,S,subjects):
    cell_list=[]
    counter=1
    from docx import Document
    from docx.shared import Inches
    document = Document()
    document.add_picture("university.png", width=Inches(0.70))
    document.add_heading(universty, 0)
    #document.add_paragraph('--------------------------------------------------------------------------------------------------------------------')
    document.add_heading('RESULT CARD', level=1)
    p=document.add_paragraph('')
    p.add_run('    NAME                                 : ').bold=True
    p.add_run(name+'\n')
    p.add_run('    REGISTRATION  NO      : ').bold=True
    p.add_run(reg_no+'\n')
    p.add_run('    DEPARTNMENT             : ').bold=True
    p.add_run(department+'\n')
    p.add_run('    DEGREE                             : ').bold=True
    p.add_run(degree+'\n')
    p.add_run('    SECTION                            : ').bold=True
    p.add_run(section+'\n')
    p.add_run('    SEMYSTER                        : ').bold=True
    p.add_run(semyster)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    table = document.add_table(rows=1, cols=7)
    cell_0 = table.cell(0, 0)
    cell_1 = table.cell(0, 1)
    cell_2 = table.cell(0, 2)
    cell_3 = table.cell(0, 3)
    cell_4 = table.cell(0, 4)
    cell_5 = table.cell(0, 5)
    cell_6 = table.cell(0, 6)
    cell_0.text='Sr#'
    cell_1.text='COURSES'
    cell_2.text='CREDIT HOURS'
    cell_3.text='TOTAL MARKS'
    cell_4.text='OBTAINED MARKS'
    cell_5.text='   %'
    cell_6.text='GRADE'
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    t = document.add_table(rows=subjects, cols=7)
    for i in range(0,subjects):
        for j in range(0,7):
            counter+=1
            cell_list.append(t.cell(i,j))

    a=0
    b=0
    for i in range(0,subjects):
        cell_list[a].text=str(sr_list[b])
        a+=1
        cell_list[a].text=subject_name[b]
        a+=1
        cell_list[a].text=str(c_h_List[b])
        a+=1
        cell_list[a].text=str(t_m_List[b])
        a+=1
        cell_list[a].text=str(obtained_marks_list[b])
        a+=1
        cell_list[a].text=str(per_list[b])
        a+=1
        cell_list[a].text=grade_list[b]
        a+=1
        b+=1

    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    t=document.add_table(rows=3,cols=4)
    cell_0 = t.cell(0, 0)
    cell_1 = t.cell(0, 1)
    cell_2 = t.cell(0, 2)
    cell_3 = t.cell(0, 3)
    cell_4 = t.cell(1, 0)
    cell_5 = t.cell(1, 1)
    cell_6 = t.cell(1, 2)
    cell_7 = t.cell(1, 3)
    cell_8 = t.cell(2, 0)
    cell_9 = t.cell(2, 1)
    cell_10 = t.cell(2, 2)
    cell_11 = t.cell(2, 3)
#GPA,T_M,O_M,p,G,S
    cell_0.text='GPA'
    cell_1.text=str(GPA)
    cell_2.text='TOTAL MARKS'
    cell_3.text=str(T_M)
    cell_4.text='OBTAINED MARKS'
    cell_5.text=str(O_M)
    cell_6.text='PERCENTAGE'
    cell_7.text=str(pe)
    cell_8.text='GRADE'
    cell_9.text=str(G)
    cell_10.text='STATUS'
    cell_11.text=str(S)
    p=document.add_paragraph('\n                                                                                                                                 By : ')
    p.add_run('SHAKAIB  HASSAN\n').bold=True
    p.add_run('                                                                                                                                               (PMAS-UIIT)')
    a=degree+'-'+semyster + ' Result card'
    b=username()
    filename="C:\\Users\\"+b+"\\Desktop\\"+a+".docx"
    #document.add_page_break()
    document.save(filename)    



clear()
universty,department,name,reg_no,degree,section,semyster=personal_info()
aa=input("\n\npress any key to continue ?")
clear()
GPA,T_M,O_M,p,G,S,subjects=result_calculator()
aa=input("\n\npress any key to continue : ")
clear()
create_table(sr_list,subject_name,c_h_List,t_m_List,obtained_marks_list,per_list,grade_list,GPA,T_M,O_M,p,G,S)
aa=input("\n\npress any key to continue : ")
clear()
choice=input("\n\n\t You want to save? (y/n) :")
if choice=='y' or choice=='Y':
    create_world_doc(universty,department,name,reg_no,degree,section,semyster,sr_list,subject_name,c_h_List,t_m_List,obtained_marks_list,per_list,grade_list,GPA,T_M,O_M,p,G,S,subjects)
    clear()
    print("Your document saved on 'DESKSTOP' with (.docx) extention.....")
    aa=input("\n\npress any key to exit ?")
else:
    print("\n\n            ------------------OK----------------------")
    aa=input("\n\npress any key to exit ?")
